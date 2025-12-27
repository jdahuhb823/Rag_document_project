
from pathlib import Path
from typing import List, Union, Optional


from langchain_core.documents import Document
from langchain_community.document_loaders import (
	PyPDFLoader,
	TextLoader,
	Docx2txtLoader,
	UnstructuredWordDocumentLoader,
	UnstructuredPDFLoader,
	UnstructuredPowerPointLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter

from bs4 import BeautifulSoup
import json
import pandas as pd
from lxml import etree


try:
	import docx as python_docx
except Exception:
	python_docx = None

try:
	import openpyxl
except Exception:
	openpyxl = None




if Document is None:
	from dataclasses import dataclass

	@dataclass
	class Document:
		page_content: str
		metadata: dict = None


class _FallbackTextSplitter:
	def __init__(self, chunk_size: int = 800, chunk_overlap: int = 100) -> None:
		self.chunk_size = int(chunk_size)
		self.chunk_overlap = int(chunk_overlap)

	def split_text(self, text: str) -> List[str]:
		if text is None:
			return []
		text = text.strip()
		if not text:
			return []

		chunks: List[str] = []
		start = 0
		L = len(text)
		while start < L:
			end = min(start + self.chunk_size, L)
			sub = text[start:end]
			sep = sub.rfind("\n\n")
			if sep != -1 and sep > int(self.chunk_size * 0.33):
				end = start + sep
			chunk = text[start:end].strip()
			if chunk:
				chunks.append(chunk)
			if end >= L:
				break
			start = max(0, end - self.chunk_overlap)
		return chunks

	def split_documents(self, docs: List[Document]) -> List[Document]:
		out: List[Document] = []
		for d in docs:
			text = getattr(d, "page_content", None) or getattr(d, "text", None) or str(d)
			md = dict(getattr(d, "metadata", {}) or {})
			for c in self.split_text(text):
				out.append(Document(page_content=c, metadata=md))
		return out


def get_text_splitter(chunk_size: int = 800, chunk_overlap: int = 100):
	if RecursiveCharacterTextSplitter is not None:
		return RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
	return _FallbackTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)



def _choose_pdf_loader():
	if PyPDFLoader is not None:
		return PyPDFLoader
	if UnstructuredPDFLoader is not None:
		return UnstructuredPDFLoader
	raise RuntimeError("No PDF loader is available. Install langchain and its PDF extras.")


def _choose_pptx_loader():
	if 'UnstructuredPowerPointLoader' in globals() and UnstructuredPowerPointLoader is not None:
		return UnstructuredPowerPointLoader
	raise RuntimeError("No PPTX loader available. Install langchain with unstructured or the pptx extras.")


def _choose_docx_loader():
	if Docx2txtLoader is not None:
		return Docx2txtLoader
	if UnstructuredWordDocumentLoader is not None:
		return UnstructuredWordDocumentLoader
	raise RuntimeError("No DOCX loader is available. Install langchain and its word/document extras.")


def _get_loader_for_path(path: Union[str, Path]):
	p = Path(path)
	if not p.exists():
		raise FileNotFoundError(f"File not found: {p}")

	suffix = p.suffix.lower()
	if suffix == ".pdf":
		LoaderCls = _choose_pdf_loader()
		def _pdf_loader():
			loader = LoaderCls(str(p))
			if hasattr(loader, "load"):
				return loader.load()
			if hasattr(loader, "load_and_split"):
				return loader.load_and_split()
			raise RuntimeError("PDF loader class does not implement 'load' or 'load_and_split'")
		return _pdf_loader
	if suffix == ".txt":
		if TextLoader is None:
			raise RuntimeError("TextLoader is not available. Ensure langchain is installed.")
		def _txt_loader():
			loader = TextLoader(str(p), encoding="utf-8")
			if hasattr(loader, "load"):
				return loader.load()
			raise RuntimeError("Text loader class does not implement 'load'")
		return _txt_loader
	if suffix == ".docx":
		# Prefer langchain-provided DOCX loaders when available
		try:
			LoaderCls = _choose_docx_loader()
			def _docx_loader():
				loader = LoaderCls(str(p))
				if hasattr(loader, "load"):
					result = loader.load()
				elif hasattr(loader, "load_and_split"):
					result = loader.load_and_split()
				else:
					raise RuntimeError("DOCX loader class does not implement 'load' or 'load_and_split'")
				return result
			return _docx_loader
		except Exception:
			# Fallback to python-docx if available
			if python_docx is None:
				raise RuntimeError("No DOCX loader is available. Install python-docx or langchain extras for DOCX support.")

			def _docx_fallback():
				doc = python_docx.Document(str(p))
				parts = []
				for para in doc.paragraphs:
					text = para.text.strip()
					if text:
						parts.append(text)
				for table in doc.tables:
					head = []
					if table.rows:
						for c in table.rows[0].cells:
							head.append(c.text.strip())
					for r in table.rows[1:]:
						cells = [c.text.strip() for c in r.cells]
						if any(cells):
							if any(head):
								pairs = [f"{h}: {v}" for h, v in zip(head, cells)]
								parts.append("; ".join(pairs))
				return "\n\n".join(parts)

			return _docx_fallback
	if suffix == ".pptx":
		LoaderCls = _choose_pptx_loader()
		def _pptx_loader():
			loader = LoaderCls(str(p))
			if hasattr(loader, "load"):
				return loader.load()
			if hasattr(loader, "load_and_split"):
				return loader.load_and_split()
			raise RuntimeError("PPTX loader class does not implement 'load' or 'load_and_split'")
		return _pptx_loader
	if suffix == ".xml":
		def _xml_loader():
			parser = etree.parse(str(p))
			root = parser.getroot()
			texts = []
			for elem in root.iter():
				if elem.text and elem.text.strip():
					texts.append(elem.text.strip())
			return "\n\n".join(texts)
		return _xml_loader
	if suffix == ".json":
		def _json_loader():
			with open(p, "r", encoding="utf-8") as fh:
				obj = json.load(fh)
			return json.dumps(obj, indent=2, ensure_ascii=False)
		return _json_loader
	if suffix == ".html" or suffix == ".htm":
		def _html_loader():
			with open(p, "r", encoding="utf-8") as fh:
				soup = BeautifulSoup(fh, "html.parser")
				for s in soup(["script", "style"]):
					s.delete()
				return soup.get_text(separator="\n")
		return _html_loader
	if suffix == ".csv":
		def _csv_loader():
			df = pd.read_csv(p)
			rows = []
			for _, r in df.iterrows():
				parts = [f"{c}: {r[c]}" for c in df.columns]
				rows.append("; ".join(parts))
			return "\n\n".join(rows)
		return _csv_loader

	if suffix in (".xls", ".xlsx"):
		def _excel_loader():
			df_dict = pd.read_excel(p, sheet_name=None)
			parts = []
			for sheet_name, df in df_dict.items():
				parts.append(f"SHEET: {sheet_name}")
				if df.empty:
					parts.append("(empty sheet)")
				else:
					for _, r in df.iterrows():
						cells = [f"{c}: {r[c]}" for c in df.columns]
						parts.append("; ".join(cells))
			return "\n\n".join(parts)
		return _excel_loader
def load_and_split_documents(
	paths: Union[str, Path, List[Union[str, Path]]],
	chunk_size: int = 800,
	chunk_overlap: int = 100,
) -> List[Document]:
	if RecursiveCharacterTextSplitter is None and _FallbackTextSplitter is None:
		raise RuntimeError(
			"RecursiveCharacterTextSplitter is not available. Ensure langchain is installed or use the fallback splitter."
		)

	if isinstance(paths, (str, Path)):
		sources = [paths]
	else:
		sources = list(paths)

	all_docs: List[Document] = []
	splitter = get_text_splitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

	for src in sources:
		p = Path(src)
		raw_texts: List[str] = []
		if isinstance(src, (str, Path)) and p.exists():
			loader = _get_loader_for_path(p)
			result = loader()
			# Normalize various loader return types into Document objects for splitting
			docs_to_split = []
			if isinstance(result, list):
				if result and isinstance(result[0], Document):
					docs_to_split = list(result)
				elif all(isinstance(x, str) for x in result):
					for t in result:
						if t and t.strip():
							docs_to_split.append(Document(page_content=t, metadata={"source": str(p)}))
				else:
					# Try to extract text-like attributes from list elements
					for item in result:
						text = getattr(item, "page_content", None) or getattr(item, "text", None) or (str(item) if isinstance(item, str) else None)
						if text and text.strip():
							docs_to_split.append(Document(page_content=text, metadata={"source": str(p)}))
			elif isinstance(result, Document):
				docs_to_split = [result]
				for d in docs_to_split:
					if d.metadata is None:
						d.metadata = {}
					d.metadata.setdefault("source", str(p))
			elif isinstance(result, str):
				text = result.strip()
				if text:
					docs_to_split = [Document(page_content=text, metadata={"source": str(p)})]
			elif result is None:
				docs_to_split = []
			else:
				# fallback: stringify
				text = str(result).strip()
				if text:
					docs_to_split = [Document(page_content=text, metadata={"source": str(p)})]

			# split and deduplicate at doc level
			if docs_to_split:
				chunks = splitter.split_documents(docs_to_split)
				# trim, remove empty, deduplicate by text
				seen = set()
				cleaned_chunks = []
				for c in chunks:
					text = getattr(c, "page_content", None) or getattr(c, "text", None) or str(c)
					if not text or not text.strip():
						continue
					k = text.strip()[:500]
					if k in seen:
						continue
					seen.add(k)
					cleaned_chunks.append(Document(page_content=text.strip(), metadata=(getattr(c, "metadata", {}) or {})))
				all_docs.extend(cleaned_chunks)
		else:
			raw_texts.append(str(src))

		for t in raw_texts:
			if not t or not t.strip():
				continue
			doc = Document(page_content=t, metadata={"source": None})
			chunks = splitter.split_documents([doc])
			all_docs.extend(chunks)

	# final cleanup: merge consecutive small chunks if needed
	final_docs: List[Document] = []
	buffer = ""
	for d in all_docs:
		text = getattr(d, "page_content", None) or getattr(d, "text", None) or str(d)
		if len(text) < 120 and buffer:
			buffer = buffer + "\n\n" + text
			if len(buffer) >= 120:
				final_docs.append(Document(page_content=buffer, metadata=getattr(d, "metadata", {}) or {}))
				buffer = ""
		else:
			if buffer:
				final_docs.append(Document(page_content=buffer, metadata=getattr(d, "metadata", {}) or {}))
				buffer = ""
			final_docs.append(d)
	if buffer:
		final_docs.append(Document(page_content=buffer, metadata={}))

	return final_docs


load_and_split = load_and_split_documents


if __name__ == "__main__":
	import argparse

	parser = argparse.ArgumentParser(description="Load and chunk documents for RAG.")
	parser.add_argument("paths", nargs="+", help="Files to load (pdf, txt, docx)")
	parser.add_argument("--chunk-size", type=int, default=800)
	parser.add_argument("--overlap", type=int, default=100)
	args = parser.parse_args()

	docs = load_and_split(args.paths, chunk_size=args.chunk_size, chunk_overlap=args.overlap)

